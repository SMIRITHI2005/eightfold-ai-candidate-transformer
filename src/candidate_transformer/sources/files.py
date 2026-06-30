"""File-based source adapters and detection helpers."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from ..models import EvidenceSourceType
from ..normalization import normalize_text
from .base import SourceAdapter, SourceDocument
from .web import WebProfileAdapter


class TextFileAdapter(SourceAdapter):
    """Adapter for text-like sources."""

    def __init__(self, source_type: EvidenceSourceType):
        self.source_type = source_type

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in {".txt", ".md", ".log", ".pdf", ".docx"}

    def load(self, path: Path) -> SourceDocument:
        text = read_text_content(path)
        return SourceDocument(
            path=path,
            source=str(path),
            source_type=self.source_type,
            text=text,
            metadata={"extension": path.suffix.lower()},
        )


class JSONFileAdapter(SourceAdapter):
    """Adapter for JSON-based sources."""

    def __init__(self, source_type: EvidenceSourceType):
        self.source_type = source_type

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".json"

    def load(self, path: Path) -> SourceDocument:
        raw = json.loads(path.read_text(encoding="utf-8"))
        text = json.dumps(raw, ensure_ascii=False, indent=2)
        return SourceDocument(
            path=path,
            source=str(path),
            source_type=self.source_type,
            text=text,
            raw=raw,
            metadata={"extension": path.suffix.lower()},
        )


class CSVFileAdapter(SourceAdapter):
    """Adapter for CSV-based structured sources."""

    def __init__(self, source_type: EvidenceSourceType):
        self.source_type = source_type

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".csv"

    def load(self, path: Path) -> SourceDocument:
        with path.open(newline="", encoding="utf-8") as handle:
            reader = csv.DictReader(handle)
            rows = [dict(row) for row in reader]

        first_row = rows[0] if rows else {}
        text = json.dumps(rows, ensure_ascii=False, indent=2)
        return SourceDocument(
            path=path,
            source=str(path),
            source_type=self.source_type,
            text=text,
            raw=first_row,
            metadata={"extension": path.suffix.lower(), "row_count": len(rows)},
        )


DEFAULT_ADAPTERS: list[SourceAdapter] = [
    CSVFileAdapter(EvidenceSourceType.ats_csv),
    JSONFileAdapter(EvidenceSourceType.ats_json),
    JSONFileAdapter(EvidenceSourceType.github_profile_json),
    JSONFileAdapter(EvidenceSourceType.linkedin_export),
    WebProfileAdapter(EvidenceSourceType.github_profile_web),
    WebProfileAdapter(EvidenceSourceType.linkedin_profile_web),
    TextFileAdapter(EvidenceSourceType.resume_pdf),
    TextFileAdapter(EvidenceSourceType.resume_docx),
    TextFileAdapter(EvidenceSourceType.resume_txt),
    TextFileAdapter(EvidenceSourceType.recruiter_notes_txt),
]


def read_text_content(path: Path) -> str:
    """Read text content from supported file formats."""

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".log"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError("Reading PDF resumes requires the optional 'pypdf' dependency") from exc
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if suffix == ".docx":
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError("Reading DOCX resumes requires the optional 'python-docx' dependency") from exc
        document = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return path.read_text(encoding="utf-8", errors="ignore")


def detect_source_type(path: Path | str, override: str | None = None) -> EvidenceSourceType:
    """Detect a source type from filename and extension."""

    path_text = str(path)
    if path_text.startswith("http://") or path_text.startswith("https://"):
        lower_url = path_text.lower()
        if "github.com" in lower_url:
            return EvidenceSourceType.github_profile_web
        if "linkedin.com" in lower_url:
            return EvidenceSourceType.linkedin_profile_web

    if override and override != "auto":
        return EvidenceSourceType(override)

    path_obj = path if isinstance(path, Path) else Path(path_text)
    name = path_obj.name.lower()
    suffix = path_obj.suffix.lower()
    if suffix == ".json":
        if "github" in name:
            return EvidenceSourceType.github_profile_json
        if "linkedin" in name:
            return EvidenceSourceType.linkedin_export
        return EvidenceSourceType.ats_json
    if suffix == ".csv":
        if "linkedin" in name:
            return EvidenceSourceType.linkedin_export
        return EvidenceSourceType.ats_csv
    if suffix == ".pdf":
        return EvidenceSourceType.resume_pdf
    if suffix == ".docx":
        return EvidenceSourceType.resume_docx
    if suffix in {".txt", ".md", ".log"}:
        if "note" in name or "recruit" in name:
            return EvidenceSourceType.recruiter_notes_txt
        if "linkedin" in name:
            return EvidenceSourceType.linkedin_export
        return EvidenceSourceType.resume_txt
    return EvidenceSourceType.resume_txt


def auto_detect_adapter(
    path: Path | str,
    source_type: str | None = None,
    adapters: list[SourceAdapter] | None = None,
) -> SourceAdapter:
    """Return the best adapter for the path."""

    resolved_type = detect_source_type(path, source_type)
    candidate_adapters = adapters or DEFAULT_ADAPTERS
    for adapter in candidate_adapters:
        if adapter.source_type == resolved_type and adapter.can_handle(path):
            return adapter
    if resolved_type in {EvidenceSourceType.ats_json, EvidenceSourceType.github_profile_json, EvidenceSourceType.linkedin_export}:
        return JSONFileAdapter(resolved_type)
    return TextFileAdapter(resolved_type)


def is_remote_source(source: Path | str) -> bool:
    """Return True when the source reference is an HTTP(S) URL."""

    text = str(source)
    return text.startswith("http://") or text.startswith("https://")


def flatten_json_payload(payload: Any, prefix: str = "") -> dict[str, Any]:
    """Flatten a nested JSON payload into dotted keys."""

    flattened: dict[str, Any] = {}
    if isinstance(payload, dict):
        for key, value in payload.items():
            nested_prefix = f"{prefix}.{key}" if prefix else str(key)
            flattened.update(flatten_json_payload(value, nested_prefix))
    elif isinstance(payload, list):
        flattened[prefix] = payload
    else:
        flattened[prefix] = payload
    return flattened


def value_as_text(value: Any) -> str:
    """Convert a JSON value to deterministic text."""

    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return normalize_text(value)
